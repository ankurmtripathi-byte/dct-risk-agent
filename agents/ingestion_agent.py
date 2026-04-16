import io
import json
import os
import sqlite3
from datetime import datetime
from pathlib import Path

import anthropic
from flask import Blueprint, jsonify, render_template, request

import config
from config import ANTHROPIC_API_KEY, DB_PATH, MODEL_ANALYSIS, UPLOAD_FOLDER
from database import get_db

client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)

ingestion_bp = Blueprint("ingestion", __name__)

ALLOWED_EXTENSIONS = {"pdf", "docx", "xlsx", "csv", "txt"}


def _allowed(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def _extract_text(filepath: str, ext: str) -> str:
    if ext == "pdf":
        try:
            import PyPDF2
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                return "\n".join(
                    page.extract_text() or "" for page in reader.pages
                )
        except Exception as e:
            return f"[PDF extraction error: {e}]"

    elif ext == "docx":
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            return f"[DOCX extraction error: {e}]"

    elif ext == "xlsx":
        try:
            import openpyxl
            wb = openpyxl.load_workbook(filepath, data_only=True)
            lines = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    line = "\t".join(str(c) if c is not None else "" for c in row)
                    if line.strip():
                        lines.append(line)
            return "\n".join(lines)
        except Exception as e:
            return f"[XLSX extraction error: {e}]"

    elif ext in ("txt", "csv"):
        try:
            with open(filepath, "r", errors="replace") as f:
                return f.read()
        except Exception as e:
            return f"[Text extraction error: {e}]"

    return "[Unsupported file type]"


# ── Routes ────────────────────────────────────────────────────────────────────

@ingestion_bp.route("/ingestion")
def ingestion_page():
    return render_template("ingestion.html")


@ingestion_bp.route("/api/documents", methods=["GET"])
def list_documents():
    conn = get_db()
    docs = [dict(r) for r in conn.execute(
        "SELECT id,filename,doc_type,upload_date,processed,extracted_risks_count,summary "
        "FROM ingested_documents ORDER BY upload_date DESC"
    ).fetchall()]
    conn.close()
    return jsonify(docs)


@ingestion_bp.route("/api/ingest", methods=["POST"])
def ingest_document():
    if "file" not in request.files:
        return jsonify(error="No file uploaded"), 400
    file = request.files["file"]
    if not file.filename or not _allowed(file.filename):
        return jsonify(error=f"Unsupported file type. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"), 400

    doc_type = request.form.get("doc_type", "risk_register")
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)

    safe_name = f"{datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{file.filename}"
    filepath  = os.path.join(UPLOAD_FOLDER, safe_name)
    file.save(filepath)

    ext  = safe_name.rsplit(".", 1)[1].lower()
    text = _extract_text(filepath, ext)
    now  = datetime.utcnow().isoformat()

    conn = get_db()
    conn.execute("""
        INSERT INTO ingested_documents
          (filename, doc_type, upload_date, processed, extracted_risks_count, extracted_text, summary)
        VALUES (?,?,?,0,0,?,?)
    """, (file.filename, doc_type, now, text[:50000], ""))
    conn.commit()
    doc_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    conn.close()

    return jsonify(id=doc_id, filename=file.filename, chars=len(text)), 201


@ingestion_bp.route("/api/documents/<int:doc_id>/process", methods=["POST"])
def process_document(doc_id):
    if not ANTHROPIC_API_KEY:
        return jsonify(error="ANTHROPIC_API_KEY not configured"), 500

    conn = get_db()
    doc  = conn.execute("SELECT * FROM ingested_documents WHERE id=?", (doc_id,)).fetchone()
    if not doc:
        conn.close()
        return jsonify(error="Document not found"), 404

    doc = dict(doc)
    text_snippet = (doc.get("extracted_text") or "")[:12000]

    prompt = f"""You are a DCT risk analyst. Analyse the following document and extract structured risk information.

Document type: {doc['doc_type']}
Filename: {doc['filename']}

--- DOCUMENT CONTENT (truncated) ---
{text_snippet}
--- END ---

1. Write a 2-3 sentence executive summary of what this document covers.
2. Identify up to 8 distinct risks mentioned or implied in the document.

Return JSON in this exact format:
{{
  "summary": "...",
  "risks": [
    {{
      "title": "...",
      "category": "...",
      "description": "...",
      "likelihood": 1-5,
      "impact": 1-5,
      "mitigation": "..."
    }}
  ]
}}"""

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model=MODEL_ANALYSIS,
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        # strip markdown fences if present
        if raw.startswith("```"):
            raw = "\n".join(raw.split("\n")[1:])
            if raw.endswith("```"):
                raw = raw[: raw.rfind("```")]
        result = json.loads(raw)

        summary  = result.get("summary", "")
        risks_ai = result.get("risks", [])

        from datetime import datetime as dt
        now = dt.utcnow().isoformat()
        from config import DB_PATH
        saved_ids = []
        for r in risks_ai:
            from agents.risk_register_agent import _next_risk_id
            rid_str = _next_risk_id("department", conn)
            l, i = int(r.get("likelihood", 2)), int(r.get("impact", 2))
            conn.execute("""
                INSERT INTO risks
                  (risk_id,level,entity_name,category,title,description,
                   likelihood,impact,risk_score,mitigation,status,source,
                   created_date,updated_date)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """, (rid_str, "department", doc["filename"],
                  r.get("category","Operational"), r.get("title",""),
                  r.get("description",""),
                  l, i, l * i, r.get("mitigation",""),
                  "Open", "Ingested", now, now))
            saved_ids.append(conn.execute("SELECT last_insert_rowid()").fetchone()[0])

        conn.execute("""
            UPDATE ingested_documents
            SET processed=1, extracted_risks_count=?, summary=?
            WHERE id=?
        """, (len(saved_ids), summary, doc_id))
        conn.commit()
        conn.close()
        return jsonify(summary=summary, risks_extracted=len(saved_ids), risk_ids=saved_ids)

    except json.JSONDecodeError as e:
        conn.close()
        return jsonify(error=f"Could not parse AI response: {e}"), 500
    except TypeError as e:
        conn.close()
        return jsonify(error=f"API key error: {e}"), 500
    except anthropic.AuthenticationError:
        conn.close()
        return jsonify(error="Invalid Anthropic API key"), 401
    except Exception as e:  # noqa: BLE001
        conn.close()
        return jsonify(error=f"Unexpected error: {e}"), 500


@ingestion_bp.route("/api/documents/<int:doc_id>", methods=["DELETE"])
def delete_document(doc_id):
    conn = get_db()
    conn.execute("DELETE FROM ingested_documents WHERE id=?", (doc_id,))
    conn.commit()
    conn.close()
    return jsonify(deleted=True)


# ── Pipeline utility functions (stubs — logic to be implemented) ──────────────

def extract_text(filepath, filetype):
    """Extract raw text from PDF, DOCX, XLSX, or TXT file.
    Returns: dict {text, page_count, error}"""
    result = {"text": "", "page_count": 0, "error": None}
    try:
        ext = filetype.lower().strip('.')

        if ext == 'pdf':
            import PyPDF2
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                result["page_count"] = len(reader.pages)
                result["text"] = "\n".join(
                    p.extract_text() or "" for p in reader.pages
                )

        elif ext == 'docx':
            from docx import Document
            doc = Document(filepath)
            result["text"] = "\n".join(p.text for p in doc.paragraphs)
            result["page_count"] = 1

        elif ext in ['xlsx', 'xls']:
            import openpyxl
            wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                lines.append(f"--- Sheet: {sheet.title} ---")
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(c) if c is not None else "" for c in row)
                    if row_text.strip(" |"):
                        lines.append(row_text)
                result["page_count"] += 1
            result["text"] = "\n".join(lines)

        elif ext == 'txt':
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                result["text"] = f.read()
            result["page_count"] = 1

        else:
            result["error"] = f"Unsupported file type: {ext}"

    except Exception as e:
        result["error"] = str(e)

    return result


def classify_document(text_sample):
    """Send first 2000 chars to Claude Haiku.
    Returns: dict {doc_type, confidence, summary, language}
    doc_type options: risk_register, policy, procedure, lesson_learned,
    incident_report, other"""
    sample = text_sample[:2000]
    try:
        response = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=300,
            messages=[{"role": "user", "content": f"""Classify this document. Return ONLY valid JSON.

Text sample:
{sample}

Return exactly:
{{"doc_type": "risk_register|policy|procedure|lesson_learned|incident_report|other",
  "confidence": 0-100,
  "summary": "One sentence describing the document",
  "language": "EN|AR|mixed"}}"""}]
        )
        text = response.content[0].text.strip()
        if text.startswith('```'):
            text = text.split('```')[1]
            if text.startswith('json'): text = text[4:]
        return json.loads(text.strip())
    except:
        return {"doc_type": "other", "confidence": 0,
                "summary": "Could not classify", "language": "EN"}


def extract_structured_data(full_text, doc_type, source_filename):
    """Extract risks or lessons from document based on doc_type.
    Returns: dict {risks: [], lessons: [], policies: [], raw_count: int}"""
    result = {"risks": [], "lessons": [], "policies": [], "raw_count": 0}

    # Truncate to avoid token limits
    text = full_text[:6000]

    if doc_type in ['risk_register', 'incident_report']:
        prompt = f"""Extract every risk or issue from this document.
Source: {source_filename}

Document:
{text}

Return ONLY a JSON array. Each item:
{{"title": "short title",
  "category": "Safety|Operational|Financial|Reputational|Compliance|Strategic|Other",
  "description": "what the risk is",
  "likelihood": null or 1-5,
  "impact": null or 1-5,
  "mitigation": "any mitigation mentioned or null",
  "owner": "any owner mentioned or null",
  "source_doc": "{source_filename}"}}

If no explicit scores found, set likelihood and impact to null.
Return empty array [] if no risks found. JSON only."""

        try:
            resp = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=2500,
                messages=[{"role": "user", "content": prompt}]
            )
            raw = resp.content[0].text.strip()
            if raw.startswith('```'):
                raw = raw.split('```')[1]
                if raw.startswith('json'): raw = raw[4:]
            result["risks"] = json.loads(raw.strip())
            result["raw_count"] = len(result["risks"])
        except Exception as e:
            result["risks"] = []

    elif doc_type == 'lesson_learned':
        prompt = f"""Extract every lesson learned from this document.
Source: {source_filename}

Document:
{text}

Return ONLY a JSON array. Each item:
{{"event_name": "event or project name",
  "event_date": "date if mentioned or null",
  "category": "Safety|Operational|Financial|Reputational|Compliance|Other",
  "lesson_title": "short title",
  "what_happened": "what occurred",
  "root_cause": "root cause if identified",
  "corrective_action": "what was done",
  "preventive_action": "what to do next time",
  "source_doc": "{source_filename}"}}

JSON only."""

        try:
            resp = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=2000,
                messages=[{"role": "user", "content": prompt}]
            )
            raw = resp.content[0].text.strip()
            if raw.startswith('```'):
                raw = raw.split('```')[1]
                if raw.startswith('json'): raw = raw[4:]
            result["lessons"] = json.loads(raw.strip())
            result["raw_count"] = len(result["lessons"])
        except:
            result["lessons"] = []

    elif doc_type in ['policy', 'procedure']:
        prompt = f"""Extract key information from this policy/procedure document.
Source: {source_filename}

Document:
{text}

Return ONLY valid JSON:
{{"policy_name": "name of the policy",
  "key_requirements": ["requirement 1", "requirement 2"],
  "compliance_risks": ["risk if not followed 1", "risk 2"],
  "relevant_roles": ["role 1", "role 2"],
  "source_doc": "{source_filename}"}}

JSON only."""

        try:
            resp = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=1000,
                messages=[{"role": "user", "content": prompt}]
            )
            raw = resp.content[0].text.strip()
            if raw.startswith('```'):
                raw = raw.split('```')[1]
                if raw.startswith('json'): raw = raw[4:]
            result["policies"] = [json.loads(raw.strip())]
            result["raw_count"] = 1
        except:
            result["policies"] = []

    return result


def reconcile_with_existing(extracted_items, db_connection):
    """Compare extracted items against DB.
    Returns: dict {new: [], duplicates: [], updates: [], report_summary: str}"""
    result = {"new": [], "duplicates": [], "updates": [], "report_summary": ""}

    if not extracted_items:
        result["report_summary"] = "No items to reconcile."
        return result

    # Fetch existing risk titles from DB
    cursor = db_connection.cursor()
    cursor.execute("SELECT risk_id, title, description FROM risks")
    existing = cursor.fetchall()
    existing_titles = [r[1].lower() for r in existing]

    if not existing_titles:
        result["new"] = extracted_items
        result["report_summary"] = f"{len(extracted_items)} new items — register is empty, all items are new."
        return result

    # Use Claude to compare
    extracted_summary = "\n".join([
        f"- {item.get('title','')}: {item.get('description','')[:100]}"
        for item in extracted_items[:20]
    ])
    existing_summary = "\n".join([
        f"- [{r[0]}] {r[1]}: {r[2][:80] if r[2] else ''}"
        for r in existing[:30]
    ])

    try:
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=1500,
            messages=[{"role": "user", "content": f"""Compare extracted items against existing register items.

EXTRACTED (from new document):
{extracted_summary}

EXISTING (in register):
{existing_summary}

Return ONLY valid JSON:
{{"new_indices": [0, 1, 2],
  "duplicate_indices": [3, 4],
  "update_indices": [5],
  "summary": "X new risks identified, Y duplicates, Z updates to existing"}}

new_indices: items not in existing register at all
duplicate_indices: items that are essentially the same as existing
update_indices: items that exist but have new/different information
Use the array index position from the EXTRACTED list."""}]
        )
        raw = resp.content[0].text.strip()
        if raw.startswith('```'):
            raw = raw.split('```')[1]
            if raw.startswith('json'): raw = raw[4:]
        mapping = json.loads(raw.strip())

        result["new"] = [extracted_items[i] for i in mapping.get("new_indices", []) if i < len(extracted_items)]
        result["duplicates"] = [extracted_items[i] for i in mapping.get("duplicate_indices", []) if i < len(extracted_items)]
        result["updates"] = [extracted_items[i] for i in mapping.get("update_indices", []) if i < len(extracted_items)]
        result["report_summary"] = mapping.get("summary", "Reconciliation complete.")

    except:
        # Fallback: mark all as new if Claude fails
        result["new"] = extracted_items
        result["report_summary"] = f"{len(extracted_items)} items — reconciliation check skipped, all marked as new."

    return result


def process_document_pipeline(filepath, db_connection):
    """Master pipeline: extract → classify → parse → reconcile.
    Returns: dict {doc_id, doc_type, summary, extracted, reconciliation}"""
    filename = Path(filepath).name
    ext = Path(filepath).suffix.strip('.')

    # Step 1: Extract text
    extracted = extract_text(filepath, ext)
    if extracted["error"]:
        return {"error": extracted["error"], "filename": filename}

    # Step 2: Classify
    classification = classify_document(extracted["text"])
    doc_type = classification["doc_type"]

    # Step 3: Extract structured data
    structured = extract_structured_data(extracted["text"], doc_type, filename)

    # Step 4: Reconcile
    items_to_reconcile = structured["risks"] or structured["lessons"]
    reconciliation = reconcile_with_existing(items_to_reconcile, db_connection)

    # Step 5: Save document record to DB
    cursor = db_connection.cursor()
    cursor.execute("""
        INSERT OR REPLACE INTO ingested_documents
        (filename, doc_type, upload_date, processed,
         extracted_risks_count, extracted_text, summary)
        VALUES (?, ?, datetime('now'), 1, ?, ?, ?)
    """, (
        filename,
        doc_type,
        structured["raw_count"],
        extracted["text"][:5000],
        classification["summary"]
    ))
    doc_id = cursor.lastrowid
    db_connection.commit()

    return {
        "doc_id": doc_id,
        "filename": filename,
        "doc_type": doc_type,
        "classification": classification,
        "extracted": structured,
        "reconciliation": reconciliation,
        "page_count": extracted["page_count"]
    }


def seed_sample_documents(uploads_folder):
    """Create sample files for demo if uploads folder is empty."""
    import os
    os.makedirs(uploads_folder, exist_ok=True)

    if os.listdir(uploads_folder):
        return  # Already has files, skip seeding

    # Sample 1: Old risk register
    risk_register_text = """DCT EVENTS DEPARTMENT — RISK REGISTER 2024

Risk ID | Category | Risk Title | Description | Likelihood | Impact | Mitigation | Owner
R-001 | Safety | Crowd Crush at Entry Gates | Large crowds during peak entry times may cause dangerous crushing at narrow entry gates | 4 | 5 | Install crowd barriers, deploy additional stewards, stagger entry times | Head of Security
R-002 | Operational | Contractor Non-Delivery | Key AV contractor may fail to deliver equipment on time due to supply chain issues | 3 | 4 | Dual supplier agreement, 48hr equipment check protocol | Events Manager
R-003 | Reputational | Social Media Incident | Negative social media coverage during event due to poor guest experience | 3 | 3 | Media monitoring team, rapid response protocol, VIP experience manager | Communications Lead
R-004 | Financial | Ticket Revenue Shortfall | Lower than projected ticket sales due to competing events in same weekend | 3 | 4 | Early bird pricing, corporate package sales, marketing campaign | Finance Manager
R-005 | Compliance | Permit Delays | Municipality permits not issued in time causing event postponement | 2 | 5 | Apply 90 days in advance, dedicated liaison officer with Abu Dhabi Municipality | Legal & Compliance
R-006 | Safety | Medical Emergency Capacity | On-site medical team insufficient for attendance of 20,000+ | 3 | 5 | SEHA partnership, 2 ambulances on standby, trained first aiders per 500 guests | Health & Safety
R-007 | Operational | Weather Disruption | Outdoor event vulnerable to sandstorm or unexpected rainfall | 4 | 3 | Weather monitoring 72hr in advance, shelter contingency plan, postponement protocol | Operations Director
R-008 | Strategic | Headline Act Cancellation | Lead performer cancels within 48hrs of event | 2 | 5 | Contractual penalty clauses, backup performer on standby, refund policy ready | CEO Office
R-009 | Operational | Parking & Transport Failure | Insufficient parking and shuttle services causing guest dissatisfaction | 4 | 3 | Pre-book parking, 3 shuttle routes, ADDC coordination | Logistics Manager
R-010 | Reputational | VIP Protocol Breach | Mishandling of VIP or government guest causing diplomatic sensitivity | 2 | 5 | Dedicated VIP team, protocol officer, pre-event briefing for all senior staff | Protocol Officer
"""

    # Sample 2: Lessons learned
    lessons_text = """LESSONS LEARNED REPORT
Event: Abu Dhabi Summer Entertainment Programme 2024
Prepared by: Events Risk Team | Date: September 2024

LESSON 1 — CROWD MANAGEMENT
Event: Opening Night Concert, Yas Island Amphitheatre, July 2024
What Happened: Entry gates opened 30 minutes late due to IT ticketing system failure.
Approximately 3,000 guests queued in 42°C heat. Two guests required medical attention for heat exhaustion.
Root Cause: Ticketing vendor performed a software update 2 hours before gates opened without notifying DCT.
Corrective Action: Ticketing system locked from changes 72 hours before any event. Manual backup ticketing deployed.
Preventive Action: All vendor system change freezes to be written into contracts. Test run mandatory 24hrs before event.

LESSON 2 — CONTRACTOR MANAGEMENT
Event: Cultural Festival, Manarat Al Saadiyat, August 2024
What Happened: Stage construction contractor mobilised 6 hours late. Event opening delayed by 2 hours.
Live broadcast window missed. Estimated reputational and contractual cost: AED 450,000.
Root Cause: Contractor had conflicting commitment at another venue. DCT had no exclusivity clause.
Corrective Action: Issued formal notice to contractor. Penalty clause invoked.
Preventive Action: All key contractors to have exclusivity clauses for event day +2 days prior.
Backup contractor list maintained for all critical services.

LESSON 3 — WEATHER CONTINGENCY
Event: Outdoor Film Screening, Corniche, February 2024
What Happened: Unexpected wind advisory (35 knots) issued 4 hours before event.
No pre-agreed decision framework existed. Team spent 2 hours debating cancellation.
Event proceeded, screen partially damaged, guest experience poor.
Root Cause: No pre-defined weather decision matrix. Escalation path unclear.
Corrective Action: Event cancelled early and guests notified. Partial refund issued.
Preventive Action: Weather decision matrix created — wind >25 knots = postpone outdoor screening.
Decision authority assigned to Operations Director. NCMS monitoring mandatory for all outdoor events.
"""

    with open(os.path.join(uploads_folder, 'sample_risk_register_2024.txt'), 'w') as f:
        f.write(risk_register_text)

    with open(os.path.join(uploads_folder, 'sample_lessons_learned_2024.txt'), 'w') as f:
        f.write(lessons_text)

    print(">> Sample documents seeded to uploads/")
